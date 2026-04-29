"""
rag_engine.py — The core RAG (Retrieval-Augmented Generation) pipeline.

This module handles:
  1. Extracting text from PDFs using PyMuPDF (with OCR fallback)
  2. Splitting text into smaller overlapping chunks
  3. Generating vector embeddings with sentence-transformers
  4. Storing and searching embeddings using FAISS
  5. Calling the Groq LLM to generate answers in multiple modes:
       - Free chat    (no document — general knowledge)
       - Document Q&A (PDF context retrieved from FAISS)
       - Web search   (DuckDuckGo context injected)
       - Image        (vision model describes uploaded images)
"""

import os
import re
import base64
import numpy as np
import fitz  # PyMuPDF
import faiss
import pytesseract
from PIL import Image
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    docx = None

# ─── Load environment variables from .env file ────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent.parent
load_dotenv(BASE_DIR / ".env")

# ─── Config ───────────────────────────────────────────────────────────────────

# LLM provider selection:
#   "groq"   — FREE tier available at https://console.groq.com  (default)
#   "xai"    — xAI API (requires paid credits)
#   "openai" — OpenAI GPT models
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "groq").lower()

# ── Groq settings (FREE) ─────────────────────────────────────────────────────
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")

# ── Vision model (Groq) ───────────────────────────────────────────────────────
# Used for image analysis — requires a vision-capable model
GROQ_VISION_MODEL = os.getenv(
    "GROQ_VISION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct"
)

# ── xAI settings (paid) ───────────────────────────────────────────────────────
XAI_API_KEY = os.getenv("XAI_API_KEY", "")
XAI_MODEL = os.getenv("XAI_MODEL", "Groq-3-mini")

# ── OpenAI settings ───────────────────────────────────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

# Chunk sizes for splitting PDF text (in characters)
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150

# How many top chunks to retrieve for each user question
TOP_K = 10

# ─── Lazy loading for embedding model ───────────────────────────────────────
_embedding_model = None


def get_embedding_model():
    """Lazy load the embedding model only when needed."""
    global _embedding_model
    if _embedding_model is None:
        print("Loading embedding model (first run may download ~90 MB)…")
        # Use smaller model for memory-constrained environments
        _embedding_model = SentenceTransformer("paraphrase-MiniLM-L3-v2")
    return _embedding_model


# ─── Initialize LLM client based on LLM_PROVIDER ─────────────────────────────
def _make_client():
    """Create the correct LLM client based on LLM_PROVIDER env var."""
    if LLM_PROVIDER == "groq":
        from groq import Groq

        if not GROQ_API_KEY:
            raise RuntimeError(
                "GROQ_API_KEY is not set in your .env file.\n"
                "Get a free key at: https://console.groq.com"
            )
        print(f"LLM provider: Groq  |  model: {GROQ_MODEL}")
        return Groq(api_key=GROQ_API_KEY)

    elif LLM_PROVIDER == "xai":
        from openai import OpenAI

        if not XAI_API_KEY:
            raise RuntimeError("XAI_API_KEY is not set in your .env file.")
        print(f"LLM provider: xAI   |  model: {XAI_MODEL}")
        return OpenAI(api_key=XAI_API_KEY, base_url="https://api.x.ai/v1")

    elif LLM_PROVIDER == "openai":
        from openai import OpenAI

        if not OPENAI_API_KEY:
            raise RuntimeError("OPENAI_API_KEY is not set in your .env file.")
        print(f"LLM provider: OpenAI  |  model: {OPENAI_MODEL}")
        return OpenAI(api_key=OPENAI_API_KEY)

    else:
        raise RuntimeError(
            f"Unknown LLM_PROVIDER '{LLM_PROVIDER}'. " "Choose: groq | xai | openai"
        )


llm_client = _make_client()


def _get_model_name() -> str:
    """Return the active model name string."""
    return {"groq": GROQ_MODEL, "xai": XAI_MODEL, "openai": OPENAI_MODEL}.get(
        LLM_PROVIDER, GROQ_MODEL
    )


# ─── PDF Text Extraction ──────────────────────────────────────────────────────


def extract_text_from_pdfs(pdf_paths: list[str]) -> list[dict]:
    """
    Extract text from one or more PDF files using PyMuPDF.

    Returns a list of page-level dicts:
        [{"page": 1, "text": "...", "source": "my_doc.pdf"}, ...]
    """
    pages = []

    for pdf_path in pdf_paths:
        filename = os.path.basename(pdf_path)

        doc = fitz.open(pdf_path)

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()

            # OCR fallback for scanned/image-based pages
            if len(text) < 30:
                try:
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img).strip()
                    if len(ocr_text) > len(text):
                        text = ocr_text
                except Exception as e:
                    print(f"OCR failed for {filename} page {page_num}: {e}")

            if len(text) < 30:
                continue

            pages.append(
                {
                    "page": page_num,
                    "text": text,
                    "source": filename,
                }
            )

        doc.close()

    return pages


# ─── DOCX Text Extraction ───────────────────────────────────────────────────


def extract_text_from_docx(file_path: str, display_name: str = "") -> list[dict]:
    """
    Extract text from a .docx Word document using python-docx.
    Falls back to raw paragraph joining if library unavailable.
    """
    source = display_name or os.path.basename(file_path)

    if docx is None:
        return [
            {
                "page": 1,
                "text": "python-docx not installed. Run: pip install python-docx",
                "source": source,
            }
        ]

    try:
        document = docx.Document(file_path)
        full_text = "\n".join(
            para.text for para in document.paragraphs if para.text.strip()
        )

        if not full_text.strip():
            return []

        # Split into pseudo-pages of ~2000 chars
        pages = []
        chunk_size = 2000
        for i, start in enumerate(range(0, len(full_text), chunk_size), 1):
            segment = full_text[start : start + chunk_size].strip()
            if segment:
                pages.append({"page": i, "text": segment, "source": source})
        return pages
    except Exception as e:
        print(f"DOCX extraction failed for {source}: {e}")
        return []


# ─── TXT / MD Text Extraction ───────────────────────────────────────────────


def extract_text_from_txt(file_path: str, display_name: str = "") -> list[dict]:
    """
    Extract text from a plain .txt or .md file.
    """
    source = display_name or os.path.basename(file_path)
    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="replace").strip()
        if not text:
            return []
        pages = []
        chunk_size = 2000
        for i, start in enumerate(range(0, len(text), chunk_size), 1):
            segment = text[start : start + chunk_size].strip()
            if segment:
                pages.append({"page": i, "text": segment, "source": source})
        return pages
    except Exception as e:
        print(f"TXT extraction failed for {source}: {e}")
        return []


# ─── Text Chunking ────────────────────────────────────────────────────────────


def split_into_chunks(pages: list[dict]) -> list[dict]:
    """
    Split each page's text into overlapping chunks.

    Returns:
        [{"chunk_id": 0, "page": 1, "source": "doc.pdf", "text": "..."}, ...]
    """
    chunks = []
    chunk_id = 0

    for page_data in pages:
        text = page_data["text"]
        page = page_data["page"]
        source = page_data["source"]

        start = 0
        while start < len(text):
            end = start + CHUNK_SIZE
            chunk = text[start:end]
            chunk = re.sub(r"\s+", " ", chunk).strip()

            if len(chunk) > 50:
                chunks.append(
                    {
                        "chunk_id": chunk_id,
                        "page": page,
                        "source": source,
                        "text": chunk,
                    }
                )
                chunk_id += 1

            start += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


# ─── FAISS Index Builder ──────────────────────────────────────────────────────


def build_faiss_index(chunks: list[dict]) -> tuple[faiss.IndexFlatL2, list[dict]]:
    """
    Generate embeddings for every chunk and store them in a FAISS index.

    Returns:
        (index, chunks) — FAISS IndexFlatL2 + original chunk list
    """
    print(f"Generating embeddings for {len(chunks)} chunks…")

    texts = [c["text"] for c in chunks]

    embeddings = (
        get_embedding_model()
        .encode(
            texts,
            show_progress_bar=False,
            convert_to_numpy=True,
        )
        .astype("float32")
    )

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)

    print(f"FAISS index built with {index.ntotal} vectors.")
    return index, chunks


# ─── Generic text extraction helper ──────────────────────────────────────────


def _extract_text_from_bytes(file_bytes: bytes, filename: str, ext: str) -> str:
    """
    Try to extract plain text from a file based on its extension.
    Supports: .txt, .md, .csv, .html, .htm, .xml, .json, .rtf, .docx, .doc
    Falls back to raw UTF-8 decode for unknown types.
    """
    try:
        if ext in {".txt", ".md", ".csv", ".rtf"}:
            return file_bytes.decode("utf-8", errors="replace")

        if ext in {".html", ".htm"}:
            import re

            html = file_bytes.decode("utf-8", errors="replace")
            html = re.sub(
                r"<style[^>]*>.*?</style>", " ", html, flags=re.DOTALL | re.IGNORECASE
            )
            html = re.sub(
                r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE
            )
            html = re.sub(r"<[^>]+>", " ", html)
            return re.sub(r"\s+", " ", html).strip()

        if ext in {".json", ".xml"}:
            return file_bytes.decode("utf-8", errors="replace")

        if ext == ".docx":
            try:
                from docx import Document as DocxDocument
                from io import BytesIO

                doc = DocxDocument(BytesIO(file_bytes))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                pass  # fall through to raw decode

    except Exception as e:
        print(f"Text extraction failed for {filename}: {e}")

    # Final fallback
    return file_bytes.decode("utf-8", errors="replace")


# ─── Context Retrieval ────────────────────────────────────────────────────────


def retrieve_context(
    query: str,
    index: faiss.IndexFlatL2,
    chunks: list[dict],
    top_k: int = TOP_K,
) -> list[dict]:
    """
    Find the most relevant chunks for a user's query using FAISS.
    """
    query_vector = (
        get_embedding_model()
        .encode(
            [query],
            convert_to_numpy=True,
        )
        .astype("float32")
    )

    distances, indices = index.search(query_vector, top_k)

    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx == -1:
            continue
        chunk = chunks[idx].copy()
        chunk["score"] = float(dist)
        results.append(chunk)

    return results


# ─── Free Chat (No Document) ──────────────────────────────────────────────────


def generate_free_answer(
    user_question: str,
    conversation_history: list[dict],
) -> dict:
    """
    Answer any question freely using the LLM's built-in knowledge.
    Includes full conversation history so the LLM can recall past exchanges.
    """
    system_prompt = (
        "You are Nexus AI, a sharp and concise AI assistant. "
        "Answer the user's question directly and accurately.\n\n"
        "Rules:\n"
        "- Use the conversation history above to answer follow-up questions and recall prior context.\n"
        "- Give well-structured answers using Markdown (headers, bullet lists, code blocks).\n"
        "- For code, always use fenced code blocks with the language name.\n"
        "- For math or logic, show your working step by step.\n"
        "- If you are unsure, say so — do not guess or invent facts.\n"
        "- Be concise when appropriate; be detailed when the question is complex."
    )

    messages = [{"role": "system", "content": system_prompt}]
    for turn in conversation_history[-20:]:
        messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": user_question})

    response = llm_client.chat.completions.create(
        model=_get_model_name(),
        messages=messages,
        temperature=0.5,
        max_tokens=3000,
    )

    return {
        "answer": response.choices[0].message.content.strip(),
        "sources": [],
        "mode": "free",
    }


# ─── Document Q&A Answer Generation ──────────────────────────────────────────


def generate_answer(
    user_question: str,
    retrieved_chunks: list[dict],
    conversation_history: list[dict],
) -> dict:
    """
    Build a prompt from retrieved PDF context and call the LLM.

    Unlike before, the LLM is now allowed to supplement document context
    with general knowledge when the document doesn't fully cover the question.

    Returns:
        {"answer": "...", "sources": [...], "mode": "document"}
    """
    if not retrieved_chunks:
        # No chunks? Fall back to free chat
        result = generate_free_answer(user_question, conversation_history)
        result["mode"] = "document"
        return result

    context_parts = []
    for chunk in retrieved_chunks:
        context_parts.append(
            f'[Source: {chunk["source"]} | Page {chunk["page"]}]\n{chunk["text"]}'
        )
    context_text = "\n\n---\n\n".join(context_parts)

    system_prompt = (
        "You are Nexus AI, an expert document analyst. "
        "Answer the user's question using ONLY the document context provided below.\n\n"
        "Rules:\n"
        "1. Base your answer strictly on the DOCUMENT CONTEXT. Do not add information from outside the document.\n"
        "2. If the document does not contain enough information to answer, say clearly: "
        "'The document does not cover this topic.' Do not guess or hallucinate.\n"
        "3. When quoting or referencing the document, mention the source filename and page number.\n"
        "4. Use the conversation history to handle follow-up questions about the same document.\n"
        "5. Format your answer clearly with Markdown (bullet points, headers, code blocks as needed).\n\n"
        f"DOCUMENT CONTEXT:\n{context_text}"
    )

    messages = [{"role": "system", "content": system_prompt}]
    for turn in conversation_history[-20:]:
        messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": user_question})

    response = llm_client.chat.completions.create(
        model=_get_model_name(),
        messages=messages,
        temperature=0.2,
        max_tokens=3000,
    )

    answer = response.choices[0].message.content.strip()

    # Deduplicate sources
    seen = set()
    sources = []
    for chunk in retrieved_chunks:
        key = (chunk["source"], chunk["page"])
        if key not in seen:
            seen.add(key)
            sources.append({"source": chunk["source"], "page": chunk["page"]})

    return {"answer": answer, "sources": sources, "mode": "document"}


# ─── Web Search Answer Generation ────────────────────────────────────────────


def generate_web_answer(
    user_question: str,
    web_results: list[dict],
    conversation_history: list[dict],
) -> dict:
    """
    Answer a question using real-time web search results as context.

    Args:
        user_question    — The user's question
        web_results      — List of {title, url, snippet} from DuckDuckGo
        conversation_history — Prior conversation turns

    Returns:
        {"answer": "...", "sources": [{title, url}, ...], "mode": "web"}
    """
    if not web_results:
        result = generate_free_answer(user_question, conversation_history)
        result["mode"] = "web"
        return result

    context_parts = []
    for i, r in enumerate(web_results, 1):
        context_parts.append(
            f"[Web Result {i}] {r['title']}\n" f"URL: {r['url']}\n" f"{r['snippet']}"
        )
    context_text = "\n\n---\n\n".join(context_parts)

    system_prompt = (
        "You are Nexus AI. Answer the user's question using the web search results below.\n\n"
        "Rules:\n"
        "1. Answer directly and factually based on the search results.\n"
        "2. Mention the source title or URL inline when you use a specific result.\n"
        "3. Do NOT add a separate 'Sources' section — sources are shown automatically in the UI.\n"
        "4. If the results don't contain enough info, say so clearly instead of guessing.\n"
        "5. Keep the answer focused, structured, and free of filler.\n\n"
        f"WEB SEARCH RESULTS:\n{context_text}"
    )

    messages = [{"role": "system", "content": system_prompt}]
    for turn in conversation_history[-10:]:
        messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": user_question})

    response = llm_client.chat.completions.create(
        model=_get_model_name(),
        messages=messages,
        temperature=0.3,
        max_tokens=2000,
    )

    answer = response.choices[0].message.content.strip()
    sources = [{"title": r["title"], "url": r["url"]} for r in web_results]

    return {"answer": answer, "sources": sources, "mode": "web"}


# ─── Image Analysis ───────────────────────────────────────────────────────────


def analyze_image(
    image_bytes: bytes,
    user_prompt: str,
    conversation_history: list[dict],
) -> dict:
    """
    Analyze an image using a vision-capable model.

    Converts image bytes to base64, sends it to the vision model,
    and returns a detailed description / analysis.

    Returns:
        {"answer": "...", "sources": [], "mode": "image"}
    """
    # Encode image to base64
    b64_image = base64.b64encode(image_bytes).decode("utf-8")

    # Detect MIME type from the first few bytes
    if image_bytes[:4] == b"\x89PNG":
        mime = "image/png"
    elif image_bytes[:3] == b"\xff\xd8\xff":
        mime = "image/jpeg"
    elif image_bytes[:4] == b"RIFF" and image_bytes[8:12] == b"WEBP":
        mime = "image/webp"
    elif image_bytes[:6] in (b"GIF87a", b"GIF89a"):
        mime = "image/gif"
    else:
        mime = "image/jpeg"  # default fallback

    # Build the vision message
    user_message_content = [
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime};base64,{b64_image}",
            },
        },
        {
            "type": "text",
            "text": (
                user_prompt
                if user_prompt.strip()
                else "Please describe and analyze this image in detail."
            ),
        },
    ]

    messages = [
        {
            "role": "system",
            "content": (
                "You are Nexus AI, an image analysis expert. "
                "Your job is to answer the user's specific question about the image. "
                "Look carefully at everything visible: objects, text, data, colors, layout, people, diagrams, charts.\n\n"
                "Rules:\n"
                "- Prioritize answering the user's exact question first.\n"
                "- If asked to describe, give a thorough structured description.\n"
                "- If text is visible in the image, extract it accurately.\n"
                "- If charts or data are visible, read the values precisely.\n"
                "- Do not guess or invent details that are not visible.\n"
                "- Format clearly using Markdown."
            ),
        }
    ]

    for turn in conversation_history[-16:]:
        messages.append({"role": turn["role"], "content": turn["content"]})

    messages.append({"role": "user", "content": user_message_content})

    try:
        # Use Groq for vision (supports llama vision models)
        if LLM_PROVIDER == "groq":
            from groq import Groq

            vision_client = Groq(api_key=GROQ_API_KEY)
            response = vision_client.chat.completions.create(
                model=GROQ_VISION_MODEL,
                messages=messages,
                temperature=0.4,
                max_tokens=2000,
            )
        else:
            # Fallback: use the regular client (works for OpenAI vision models)
            response = llm_client.chat.completions.create(
                model=_get_model_name(),
                messages=messages,
                temperature=0.4,
                max_tokens=2000,
            )

        answer = response.choices[0].message.content.strip()
        return {"answer": answer, "sources": [], "mode": "image"}

    except Exception as e:
        return {
            "answer": f"⚠️ Image analysis failed: {str(e)}\n\nThis may be because the current model doesn't support vision. Try switching to a vision-capable model.",
            "sources": [],
            "mode": "image",
        }


# ─── URL Summarization ────────────────────────────────────────────────────────


def summarize_url_content(
    url: str,
    page_text: str,
    conversation_history: list[dict],
) -> dict:
    """
    Summarize the content of a fetched URL.

    Returns:
        {"answer": "...", "sources": [{"title": "URL", "url": url}], "mode": "url"}
    """
    system_prompt = (
        "You are Nexus AI. Provide a concise, structured summary of the web page content below.\n\n"
        "Rules:\n"
        "1. Extract only information present in the page text — do not add outside knowledge.\n"
        "2. Start with 1-2 sentences on what the page is about.\n"
        "3. Follow with a bullet list of key points and takeaways.\n"
        "4. If the content is unclear or sparse, state that briefly.\n"
        "5. Format using Markdown.\n\n"
        f"PAGE URL: {url}\n\nPAGE CONTENT:\n{page_text[:4000]}"
    )

    messages = [{"role": "system", "content": system_prompt}]
    messages.append({"role": "user", "content": "Summarize this page."})

    response = llm_client.chat.completions.create(
        model=_get_model_name(),
        messages=messages,
        temperature=0.2,
        max_tokens=1500,
    )

    answer = response.choices[0].message.content.strip()
    return {"answer": answer, "sources": [{"title": url, "url": url}], "mode": "url"}


# ─── Universal Document Analysis ──────────────────────────────────────────────


def analyze_document_with_context(
    file_bytes: bytes,
    filename: str,
    ext: str,
    question: str,
    conversation_history: list[dict],
    state,
) -> dict:
    """
    Handle any non-PDF, non-image document (Word, text, CSV, HTML, JSON, etc.).

    Strategy:
    - Extract text from the file.
    - If text is long (>300 chars), chunk + index into FAISS and do RAG.
    - If text is short, send it directly as context to the LLM.

    Returns:
        {"answer": "...", "sources": [...], "mode": "document"}
    """
    raw_text = _extract_text_from_bytes(file_bytes, filename, ext)

    if not raw_text or len(raw_text.strip()) < 20:
        return {
            "answer": f"⚠️ Could not extract readable text from **{filename}**. "
            "The file may be binary or in an unsupported format.",
            "sources": [],
            "mode": "document",
        }

    if len(raw_text) > 300:
        # Chunk and index
        pages = [{"page": 1, "text": raw_text, "source": filename}]
        chunks = split_into_chunks(pages)

        if chunks:
            all_chunks = state.chunks + chunks
            index, all_chunks = build_faiss_index(all_chunks)
            state.faiss_index = index
            state.chunks = all_chunks
            state.mode = "document"

            # Answer the user's question using retrieved context
            retrieved = retrieve_context(question, state.faiss_index, state.chunks)
            return generate_answer(question, retrieved, conversation_history)

    # Short content — inject directly into prompt
    system_prompt = (
        "You are an expert AI assistant. The user has uploaded a document whose full "
        "content is provided below. Answer the user's question based on this content.\n\n"
        f"DOCUMENT ({filename}):\n{raw_text[:4000]}\n\n"
        "Guidelines:\n"
        "1. Answer primarily from the document content.\n"
        "2. If the document doesn't cover the question, say so clearly.\n"
        "3. Use Markdown formatting for clarity."
    )

    messages = [{"role": "system", "content": system_prompt}]
    for turn in conversation_history[-20:]:
        messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": question})

    response = llm_client.chat.completions.create(
        model=_get_model_name(),
        messages=messages,
        temperature=0.3,
        max_tokens=2000,
    )

    answer = response.choices[0].message.content.strip()
    return {
        "answer": answer,
        "sources": [{"source": filename, "page": 1}],
        "mode": "document",
    }
